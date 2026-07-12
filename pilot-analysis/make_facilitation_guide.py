from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ---- Base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

ACCENT = RGBColor(0x1F, 0x3B, 0x73)   # deep blue
GREY = RGBColor(0x55, 0x55, 0x55)
STAGE = RGBColor(0x8A, 0x52, 0x00)    # brown for stage directions


def title(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = ACCENT
    p.space_after = Pt(2)
    return p


def subtitle(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = GREY
    return p


def block_heading(text):
    p = doc.add_paragraph()
    p.space_before = Pt(14)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = ACCENT
    return p


def sub_heading(text):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x2E, 0x2E, 0x2E)
    return p


def stage(text):
    """Italic stage direction (what to do, not say)."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = STAGE
    return p


def say(text):
    """A spoken script paragraph, indented and quoted feel."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def divider():
    p = doc.add_paragraph()
    r = p.add_run("\u2500" * 40)
    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ============ HEADER ============
title("Day 1 Session \u2014 Full Facilitation Guide")
subtitle("Head Tracking for Accessible Computing  |  Tue, June 16  |  9:00 AM\u201312:00 PM  |  Virtual")
divider()

# ============ BEFORE ============
block_heading("Before the session starts (do this at 8:45)")
bullet("Open the head-control app and MediaPipe demo in tabs, ready to share.")
bullet("Have a recorded backup clip of you using the system, in case the live demo fails.")
bullet("Open your follow-up email draft so you can send it the moment you finish.")
bullet("Put on camera, smile, have water nearby.")
stage("As people join, greet each by name in the chat or out loud:")
say("\u201cHi [name], welcome! Glad you made it. We\u2019ll start in just a couple minutes once everyone\u2019s here.\u201d")

# ============ BLOCK 1 ============
block_heading("BLOCK 1 \u2014 Welcome & Icebreakers (9:00 \u2013 9:30)")

sub_heading("Opening (9:00 \u2013 9:05)")
stage("Wait until everyone has joined, then start warmly.")
say("\u201cGood morning everyone, and welcome! I\u2019m so happy you\u2019re all here. My name is Sabiha Tahsin Soha \u2014 but please, just call me Soha. I want to start by saying congratulations: out of everyone who could be here, you\u2019re the ones who showed up curious about building technology that actually helps people. That\u2019s exactly the right reason to be here.")
say("Today is our very first day, so I\u2019m going to keep it relaxed. We\u2019re going to get to know each other, I\u2019ll tell you what this project is all about, and \u2014 the fun part \u2014 you\u2019re going to try the system yourselves and control a cursor using just your head. By the end of today you\u2019ll already feel like part of the research team.")
say("Before anything else: a quick promise. You do not need any coding or technical background to succeed here. None. If you have it, great. If you don\u2019t, also great \u2014 I will walk you through everything step by step. Your only job is to stay curious and ask questions. Sound good?\u201d")

sub_heading("Icebreaker (9:05 \u2013 9:25)")
say("\u201cLet\u2019s get to know each other. I\u2019ll go around one by one \u2014 no pressure, this is the easy stuff. When it\u2019s your turn, tell us three things:\u201d")
bullet("Your name and where you\u2019re joining from / what school you\u2019re at")
bullet("One thing you\u2019re excited or curious about for this summer")
bullet("And a fun one: think of an app or piece of technology you use every day \u2014 what\u2019s something about it you think is really well designed, or really badly designed?")
stage("Let each intern answer. After each, respond briefly and personally \u2014 \u201cOh nice, I love that,\u201d or \u201cThat\u2019s such a good example.\u201d")
stage("Secret purpose: that third question gets them thinking like HCI researchers \u2014 about usability and design \u2014 without you having to lecture. When someone mentions something frustrating, you can plant a seed:")
say("\u201cThat\u2019s a perfect example \u2014 that frustration you felt? That\u2019s exactly the kind of thing we study. When a tool is hard to use, it\u2019s usually not the user\u2019s fault, it\u2019s the design. Hold onto that thought, because that\u2019s the heart of what we do this summer.\u201d")

sub_heading("Your introduction (9:25 \u2013 9:30)")
say("\u201cNow a little about me, so you know who you\u2019re working with. I\u2019m a PhD student in Computer Science and Engineering at UC Santa Cruz, in the Computer Vision Lab. My research is about making webcam-based head-pointer systems work better for people with upper-limb motor impairments.")
say("Here\u2019s why I care about this. Most of us point and click without thinking about it. But for people with conditions like cerebral palsy, a spinal cord injury, or multiple sclerosis, using a mouse can be difficult or impossible. Head-pointing lets them control a computer using only head movement and a normal webcam \u2014 no expensive special hardware. The problem is that current systems are slow and clunky. My work is about making them faster, more accurate, and more personal. To me, this isn\u2019t just a technical problem \u2014 it\u2019s about independence and dignity. And this summer, you\u2019ll be helping with real research on it.")
say("One more thing: I was once exactly where you are \u2014 nervous, not sure if I belonged, no idea what half the words meant. So please never feel embarrassed to ask anything. Okay \u2014 let\u2019s talk about the project.\u201d")

# ============ BLOCK 2 ============
block_heading("BLOCK 2 \u2014 The \u201cWhy\u201d + Project Overview (9:30 \u2013 10:05)")

sub_heading("Start with empathy (9:30 \u2013 9:40)")
say("\u201cBefore I explain the technology, I want you to imagine something. Picture this: you wake up tomorrow and you can\u2019t use your hands or arms to control a computer \u2014 but your mind is sharp and you have things you want to do: message friends, do schoolwork, watch videos, browse the web. How would you use a computer?\u201d")
stage("Let them brainstorm. Call on people gently: \u201cWhat do you think, [name]?\u201d Accept all answers \u2014 voice control, eye tracking, etc. Validate each.")
say("\u201cThese are all great ideas, and they\u2019re all real technologies people use. Voice has limits in noisy rooms or for people with speech differences. Eye tracking can be tiring and expensive. The one we work on is head-based pointing \u2014 you move your head, and the cursor follows. A regular webcam watches your head, and software translates that movement into cursor movement. It\u2019s affordable because almost every laptop already has a webcam.\u201d")

sub_heading("The project in plain language (9:40 \u2013 9:55)")
say("\u201cSo here\u2019s the actual research problem. Head-pointing already exists \u2014 but most systems are slow and frustrating. Why? Because they use the same settings for everyone. But people move their heads very differently. Some people make big smooth movements, some make tiny shaky ones. A one-size-fits-all system fails a lot of people.")
say("Our big question this summer is: how do different system settings change how fast and how accurately someone can point? We look at things like:\u201d")
bullet("Smoothing \u2014 how much we \u2018calm down\u2019 shaky movement. Too little and the cursor is jittery; too much and it feels laggy. There\u2019s a sweet spot.")
bullet("Calibration \u2014 how the system learns your particular head movements before you start.")
bullet("Selection timing \u2014 how the system decides you meant to \u2018click\u2019 on something.")
say("\u201cAnd to measure all this scientifically, we use something called Fitts\u2019 Law. Don\u2019t worry about the details yet \u2014 just know it\u2019s a standard, respected method in our field for measuring how good a pointing system is. You\u2019ll genuinely become comfortable with it over the summer, I promise.\u201d")

sub_heading("Preview the summer (9:55 \u2013 10:05)")
stage("Share your screen showing the week-by-week schedule, or just talk through it.")
say("\u201cLet me show you the journey we\u2019re going on, so today makes sense in the bigger picture. Think of it as a story in chapters:\u201d")
bullet("Weeks 1\u20132: We get oriented \u2014 read a couple of beginner-friendly articles, explore the system, and start learning Python basics.")
bullet("Weeks 2\u20133: We dig into calibration and smoothing filters \u2014 testing settings and describing how the cursor feels.")
bullet("Weeks 4\u20136: We learn Fitts\u2019 Law properly, run a real pilot experiment together, and start analyzing data in Python \u2014 making charts and computing real performance metrics.")
bullet("Weeks 7\u20138: We turn everything into a presentation, and on August 8th \u2014 Presentation Day \u2014 you present your findings.")
say("\u201cSo by the end, you\u2019ll have read real research, written Python code, run an experiment, analyzed data, and presented like scientists. That\u2019s a lot \u2014 but we build up to it gently, one small step at a time. You\u2019ll be amazed how much you know in eight weeks.\u201d")

# ============ BREAK ============
block_heading("BREAK (10:05 \u2013 10:15)")
say("\u201cLet\u2019s take a 10-minute break. Grab water, stretch, step away from the screen. When we come back, the fun part: you\u2019re going to try the system yourself. Make sure you\u2019re on Google Chrome and that your webcam works when we return. Back at 10:15!\u201d")

# ============ BLOCK 3 ============
block_heading("BLOCK 3 \u2014 Hands-On: Try the System (10:15 \u2013 11:10)")

sub_heading("You demo first (10:15 \u2013 10:25)")
stage("Share your screen. Go to https://head-control-website.vercel.app.")
say("\u201cOkay, watch my screen first, then you\u2019ll do it yourself. I\u2019m opening our head-tracking website. The first step is always calibration \u2014 this is where the system learns how I move my head. I\u2019ll follow the prompts\u2026 (do it, narrating) \u2026and now watch \u2014 I\u2019m not touching my mouse at all. I\u2019m moving the cursor with just my head.\u201d")
stage("Move the cursor around, hover over things.")
say("\u201cNotice a few things as I do this: see how the cursor moves when I turn my head? Watch whether it feels smooth or a little jumpy. These are exactly the things I want you to pay attention to when it\u2019s your turn.\u201d")

sub_heading("They try it (10:25 \u2013 10:50)")
say("\u201cNow it\u2019s your turn. Everyone open Chrome and go to the same link \u2014 I\u2019ll drop it in the chat. Start with calibration, then just play with moving the cursor around. There\u2019s no way to break anything, so explore freely.")
say("As you go, I want you to be a researcher, not just a user. Keep a note open \u2014 a Google Doc or even paper \u2014 and jot down observations. Here\u2019s what to look for:\u201d")
bullet("What felt easy? What felt frustrating?")
bullet("Was the cursor smooth or jumpy/shaky?")
bullet("Was it hard or easy to land on a small target?")
bullet("What would make this better for someone who used it all day?")
say("\u201cTake about 20 minutes. I\u2019ll stay here \u2014 unmute or drop a message in the chat anytime you get stuck or notice something interesting.\u201d")
stage("Float around. Watch the chat. Help with webcam permission issues. Encourage: \u201cThat\u2019s a great observation, write that down!\u201d")

sub_heading("MediaPipe demo (10:50 \u2013 11:05)")
say("\u201cNow I want to show you what\u2019s happening under the hood. Open this next link \u2014 I\u2019ll paste it \u2014 the MediaPipe Face Landmarker demo. Allow it to use your webcam.\u201d")
stage("Paste https://mediapipe-studio.webapps.google.com/demo/face_landmarker")
say("\u201cSee all those little dots mapped onto your face? Those are called landmarks \u2014 points the computer tracks: corners of your eyes, your nose, your jaw. Move your head around and watch them follow you. This is how the computer \u2018sees\u2019 your head move. Our system takes these points and turns them into cursor movement.")
say("Things to notice: What happens when you move fast vs. slow? What happens if you tilt your head, or if the lighting is bad, or you partially cover your face? Try it \u2014 this tells you a lot about why head-tracking is hard to get right.\u201d")

sub_heading("Share-out (11:05 \u2013 11:10)")
say("\u201cLet\u2019s hear from everyone \u2014 just one observation each. What\u2019s one thing you noticed when you used the system or the landmark demo?\u201d")
stage("Go around. Affirm each: \u201cYes \u2014 that jumpiness you felt? That\u2019s exactly the smoothing problem we\u2019ll study in Week 3. You\u2019re already noticing the real research questions.\u201d")

# ============ BLOCK 4 ============
block_heading("BLOCK 4 \u2014 Logistics, Expectations & Week 1 Tasks (11:10 \u2013 11:45)")

sub_heading("Logistics (11:10 \u2013 11:25)")
say("\u201cLet\u2019s cover some quick practical things so you always know what to expect.")
say("Schedule: We meet Tuesday, Wednesday, and Thursday, 9 AM to noon. Please come to all of them \u2014 consistency is how we build momentum.")
say("How to reach me: Email is best \u2014 ssoha@ucsc.edu. I check email Monday through Friday, 9 to 5, and I\u2019ll always reply within 24 hours. I also have office hours Thursdays 8 to 9 AM by appointment if you want one-on-one time. And of course, you can ask questions anytime during our meetings.")
say("The most important rule: there is no such thing as a dumb question. If you\u2019re confused, I promise at least one other person is too, and you\u2019ll be doing them a favor by asking. I would much rather you ask than stay stuck and frustrated.")
say("If you fall behind: life happens \u2014 that\u2019s completely okay. Just tell me in advance if you can\u2019t finish something on time, and we\u2019ll figure it out. I\u2019m flexible as long as you keep me in the loop. The only thing I ask is that you don\u2019t disappear silently.")
say("One serious note \u2014 using AI tools. You\u2019re allowed to use tools like ChatGPT or Claude to help you understand a concept or fix a Python error. But two rules: first, your final work has to reflect your own understanding \u2014 not just copy-pasted AI text. Second, if you use AI for something, just say so \u2014 disclose it. Honesty here protects you and the integrity of our work. When in doubt, ask me.")
say("And quickly on confidentiality: the papers, data, and materials we use are for your learning here \u2014 please don\u2019t share or post them publicly. But anything you create \u2014 your notes, your code, your reflections \u2014 that\u2019s yours to keep and use however you want.\u201d")

sub_heading("Week 1 tasks (11:25 \u2013 11:40)")
say("\u201cHere\u2019s what to work on this first week. I\u2019ll email all of this with the links right after we finish, so don\u2019t worry about writing it all down now. None of it is meant to be stressful \u2014 it\u2019s about easing in.")
say("Reading & reflection:\u201d")
bullet("Read a short, beginner-friendly article called \u2018What is a Head Mouse System?\u2019 and write one paragraph on what you learned.")
bullet("Read \u2018What is Fitts\u2019 Law?\u2019 \u2014 it\u2019s a simple visual intro \u2014 and write one paragraph on that too.")
bullet("Spend time on the head-tracking app like we did today, and write down your observations.")
bullet("Try the MediaPipe demo again and note what you observe about how it detects your face.")
bullet("And a slightly bigger one: a roughly one-page reflection answering \u2014 What is head-based pointing? Who benefits from it, and why does it matter? Honestly, after today\u2019s discussion, you already have most of what you need for this.")
say("\u201cGetting your tools ready:\u201d")
bullet("Start the free Python tutorial at learnpython.org \u2014 just the basics: variables, data types, loops. Go at your own pace.")
bullet("Install Python 3, Visual Studio Code, and three libraries: numpy, pandas, and matplotlib.")
say("\u201cI know that setup part can feel intimidating if you\u2019ve never done it. So here\u2019s my offer: let\u2019s do a setup walkthrough together on Wednesday or Thursday, screen-sharing, step by step, so nobody gets stuck alone. Does that sound helpful?\u201d")
stage("Let them respond, pick a day.")

sub_heading("Set tone for the reflection (11:40 \u2013 11:45)")
say("\u201cFor all the writing tasks \u2014 don\u2019t stress about sounding fancy or academic. I want your honest thoughts in your own words. \u2018The cursor felt jumpy and it annoyed me\u2019 is a perfectly good research observation. Write like you\u2019re explaining it to a friend.\u201d")

# ============ BLOCK 5 ============
block_heading("BLOCK 5 \u2014 Q&A & Warm Close (11:45 \u2013 12:00)")

sub_heading("Open Q&A (11:45 \u2013 11:55)")
say("\u201cWe\u2019ve got time for questions. Anything at all \u2014 about the project, the schedule, the tech, what to expect, anything that\u2019s on your mind. Who wants to start?\u201d")
stage("If silence, prompt gently:")
say("\u201cLet me start you off with a common one people are too shy to ask: \u2018What if I\u2019m really bad at coding?\u2019 And my answer is \u2014 you won\u2019t be \u2018bad,\u2019 you\u2019ll be learning, and that\u2019s the whole point. Nobody here expects you to already know it. What else is on your mind?\u201d")
stage("If a question goes too deep/technical, use the parking lot:")
say("\u201cOoh, that\u2019s a fantastic question \u2014 and honestly it\u2019s exactly what we dig into in a few weeks. Let me write it down so we make sure to come back to it. Hold that thought.\u201d")

sub_heading("Warm close (11:55 \u2013 12:00)")
say("\u201cBefore we wrap \u2014 I want to ask you one thing: is there anything you\u2019re hoping to get out of this summer that I haven\u2019t mentioned? Anything you\u2019d love to learn or try?\u201d")
stage("Listen and acknowledge. Then close:")
say("\u201cI just want to say \u2014 I\u2019m genuinely excited to work with all of you. Remember: curiosity matters more than experience here, and I\u2019ve got your back the whole way. You did great today.")
say("Quick recap of next steps: check your email later today for the task list and all the links. Start poking at the readings and the Python tutorial when you can, and we\u2019ll do the software setup together this week. Our next meeting is Wednesday at 9 AM \u2014 same place.")
say("Thank you all so much. Have a wonderful rest of your day, and I\u2019ll see you Wednesday!\u201d")

# ============ AFTER ============
block_heading("After the session")
bullet("Send the follow-up email immediately while it\u2019s fresh: warm recap, Week 1 task list, all links, your contact info, and the day you picked for the setup walkthrough.")
bullet("Jot down a few notes about each intern (interests, comfort level, any access needs) so you can support them personally.")

out = "/Users/soha/head-control-website-js/pilot-analysis/Day1_Facilitation_Guide.docx"
doc.save(out)
print("Saved:", out)
